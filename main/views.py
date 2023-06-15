from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from .models import *
from django.contrib.auth.decorators import login_required
# Create your views here.

@login_required
def create_docx(request):
    if request.method == 'POST':
        soni = request.POST.get('soni', '')
        sana = request.POST.get('sana', '')
        xodim = request.POST.get('xodim', '')
        korxona = request.POST.get('korxona', '')
        lavozim = request.POST.get('lavozim', '')
        hissa = request.POST.get('hissa', '')
        asosiy = request.POST.get('asosiy', '')
        muayan = request.POST.get('muayan', '')
        boshlash = request.POST.get('boshlash', '')
        tomom = request.POST.get('tomom', '')
        sinov = request.POST.get('sinov', '')
        ish = request.POST.get('ish', '')
        kun = request.POST.get('kun', '')
        soat = request.POST.get('soat', '')
        razryad = request.POST.get('razryad', '')
        kalemdar1 = request.POST.get('kalemdar1', '')
        kalemdar2 = request.POST.get('kalemdar2', '')

        Xodimlar.objects.create(shartnoma_soni=soni,sana=sana,xodim=xodim,
                                korxona=korxona,lavozim=lavozim,hissa=hissa,asosiy_ish=asosiy,muddatli_mehnat=muayan,
                                ish_boshlash=boshlash,tomom_bolishi=tomom,sinov_muddati=sinov,ish_vaqti=ish,
                                kun=kun,soat=soat,razryad=razryad,kalemdar1=kalemdar1,kalemdar2=kalemdar2,
                                author=request.user)

        document = Document()

        section = document.sections[-1]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(0.5)

        style = document.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        document.add_paragraph(f'{soni}- SON MEHNAT ShARTNOMASI  ').alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('')
        p = document.add_paragraph(f'               ')
        p.add_run(f'{sana}').bold = True
        p.add_run(f'                                                                                             ')
        p.add_run('Buxoro shahri').bold = True
        a = document.add_paragraph("1. Korxona (mulkchilikning barcha shakllaridagi tashkilot, muassasa, shu jumladan, ularning alohida tarkibiy bo‘linmalari ")
        a.add_run('Buxoro davlat pedagogika instituti rektori ').bold = True
        a.add_run(f'(direktori) ')
        a.add_run(f'Daminov Mirzohid Islomovich ').bold = True
        a.add_run('nomidan, keyingi o‘rinlarda «Ish beruvchi» deb ataladi va fuqaro ')
        a.add_run(f'{xodim} ').bold = True
        a.add_run('keyingi o‘rinlarda «Xodim» deb ataladi, mazkur shartnomani quyidagilar haqida tuzdik:')
        a.name = 'Times New Roman'
        a.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('2. Xodim                                              ')
        p.add_run(f"{xodim}").bold = True
        p = document.add_paragraph('')
        p.add_run("(familiyasi, ismi va otasining ismi)").font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('')
        p.add_run(f'{korxona}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('')
        p.add_run('(korxona tarkibiy bo‘linmasi, sex, bo‘lim, uchastka, laboratoriya va shu kabilarning nomi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f"kasbi bo‘yicha ")
        p.add_run(f'{lavozim} ').bold = True
        p.add_run('lavozimiga ')
        p.add_run(f'{hissa} ').bold = True
        p.add_run('hissa ishga qabul qilinadi.')
        p = document.add_paragraph('')
        p.add_run('(ХALIKK bo‘yicha kasb, lavozimining to‘liq nomi, razryad, malaka toifasi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f'3. Shartnoma ')
        p.add_run(f'{asosiy} ').bold = True
        p.add_run('hisoblanadi. ')

        p = document.add_paragraph('')
        p.add_run('(asosiy ish, o‘rindoshlik, mavsumiy va boshqa)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f'4. Shartnoma muddati ')
        p.add_run(f'{muayan} ').bold = True

        p = document.add_paragraph('')
        p.add_run('(nomuayyan muddatga,uch yildan ko‘p bo‘lmagan muayyan muddatga (muddatli mehnat shartnomasi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph('5. Shartnoma bo‘yicha ishlash')
        p = document.add_paragraph(f'– boshlanishi ')
        p.add_run(f'{boshlash} ').bold = True


        p = document.add_paragraph(f'– тamom bo‘lishi ')
        p.add_run(f'{tomom}').bold = True

        p = document.add_paragraph(f'6. Sinov muddati ')
        p.add_run(f'{sinov}').bold = True


        p = document.add_paragraph('')
        p.add_run('(sinovsiz, sinov muddati)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('7. ')
        p.add_run('Xodimning majburiyatlari:').bold = True

        document.add_paragraph("a) O‘zbekiston Respublikasi Konstitutsiyasi, mehnat to‘g‘risidagi qonunchilik hamda boshqa qonunlar va qonunchilik hujjatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph('b) Lavozim yo‘riqnomasida belgilangan o‘z xizmat majburiyatlarini vijdonan bajarishi,').alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("d) Ish beruvchi, yuqori turuvchi organ va tashkilotlarning buyruqlari, topshiriqlari, qarorlari, ko‘rsatmalarini bajarish (bunday buyruq, topshiriq, qaror, ko‘rsatma yaqqol qonunchilik hujjatlariga zid bo‘lgan holatlar bundan mustasno);").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("e) mehnat va texnologiya intizomi (Ustav, Ichki mehnat tartibi qoidalari, Odob-axloq kodeksi va boshqa lokal hujjatlar)ga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("f) ish beruvchining qonuniy talablarini bajarish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("g) mehnatni muhofaza qilish, xavfsizlik texnikasi va ishlab chiqarish sanitariyasi talablariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("h) lavozim yo‘riqnomalariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("i) XALIKK, MM bilan nazarda tutilgan malaka majburiyatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("j) qonunchilik va boshqa normativ hujjatlarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("k) jamoa shartnomasi shartlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("l) institutning korrupsiyaga qarshi kurashish siyosati va sohaga oid boshqa normativ-huquqiy hujjatlarda qoʻyilgan maqsad va vazifalarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("m) oʻzi toʻgʻrisidagi shaxsiy ma’lumotlarning haqqoniyligini zimmasiga olish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("n) Ish beruvchining ruxsatisiz oshkor qilib bo‘lmaydigan (konfidensial) ma’lumotlarni o‘z ichiga oluvchi axborotlarni uchinchi shaxsga bermaslik, shuningdek, shaxsga oid ma’lumotlarni va xizmat doirasida foydalanilgan axborotlarni sir saqlanishiga rioya etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("o) institut professor-o‘qituvchilari va boshqa toifadagi xodimlar ish vaqti davomida o‘zaro munosabatlarda umume’tirof etilgan axloq qoidalariga rioya etish va boshqa xodim va talabalarga nisbatan xushmuomalada bo‘lish.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('8. ')
        p.add_run('Ish beruvchining majburiyatlari:').bold = True

        document.add_paragraph("a) O‘zbekiston Respublikasi Konstitutsiyasi, mehnat to‘g‘risidagi qonunchilik hamda boshqa qonunlar va qonunchilik hujjatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("b) ishga qabul qilish chogʻida xodimni Ustav, Ichki mehnat tartib-qoidalari, Odob-axloq kodeksi, mehnatini tashkil etish, mehnatni muhofaza qilish va texnik xavfsizlikni ta’minlash qoidalari, Lavozim yo‘riqnomasi, jamoa shartnomasi va boshqa ichki lokal hujjatlar bilan tanishtirish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("d) mehnat va ishlab chiqarish intizomini ta’minlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("e) ish haqi va qonunchilik hujjatlarida nazarda tutilgan boshqa to‘lovlarni o‘z vaqtida hamda to‘liq hajmda to‘lash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("f) xizmat vazifalarini bajarishda zarur xavfsiz va samarali mehnat sharoitlarini yaratishni ta’minlash, uni kasbiy kompetensiyalarini oshirish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("g) ish joyini mehnatni muhofaza qilish va xavfsizlik texnikasi qoidalariga muvofiq jihozlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("h) qonunchilikka va boshqa normativ hujjatlarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("i) jamoa shartnomasi shartlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("j) ish beruvchi tomonidan qabul qilinadigan boshqa majburiyatlar:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("k) Ish beruvchi qonunchilik hujjatlarida belgilangan boshqa majburiyatlarni ham bajarishi mumkin.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("9. Mehnat shartnomasini tuzish va bekor qilish, shuningdek unga o‘zgartish va qo‘shimchalar kiritish qonunchilikda belgilangan tartibda «Yagona milliy mehnat tizimi» idoralararo dasturiy-apparat kompleksida ro‘yxatdan o‘tkazilishi shart.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph("10. Ish kuni rejimi ")
        p.add_run(f"{ish} ").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph('')
        p.add_run('(ish vaqtining normal davomiyligi; ish vaqtining qisqartirilgan davomiyligi; to‘liqsiz ish vaqti.').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f"11.Xodim uchun ")
        p.add_run(f'{kun} ').bold = True
        p.add_run('kunlik ish haftasi belgilanadi. Bunda ish vaqtining davomiyligi haftasiga ')
        p.add_run(f'{soat} ').bold = True
        p.add_run('soatgacha etib belgilanadi.')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("12. Mehnatga haq to‘lash.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("Xodimga quyidagicha haq to‘lash belgilanadi:")
        p = document.add_paragraph(f"a) shtatlar jadvaliga muvofiq ")
        p.add_run(f'{razryad} ').bold = True
        p.add_run(f"asosida")
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('(to‘lov turi va uning aniq summadagi yoki YTS razryadi ko‘rsatilgan holdagi, yohud tushumdan olingan foizlardagi miqdori)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph("b) amaldagi qonunchilikka va normativ hujjatlarga muvofiq mehnat sharoitlari bilan bog‘liq bo‘lgan qo‘shimcha haq, ustama, kompensatsiyalar quyidagi miqdorlarda:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('        (qo‘shimcha haq, ustama, kompensatsiyalar nomi)                                            (ularning miqdori) ').font.size = Pt(10)

        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("v) jamoa shartnomasi shartlari bilan nazarda tutilgan, shuningdek, berilgan (mavjud) huquqlar va mablag‘lar doirasida rahbar tomonidan belgilanadigan qo‘shimcha haq, ustama, mukofot, taqdirlashlar va rag‘batlantiruvchi turdagi boshqa to‘lovlar;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('        (qo‘shimcha haq, ustama, kompensatsiyalar nomi)                                            (ularning miqdori) ').font.size = Pt(10)
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("13. Xodimga:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("a) asosiy ta’til (mehnat ta’tili) 21 kalendar kun;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph("b) jamoa shartnomasi boʻyicha qo‘shimcha ta’til ")
        p.add_run(f'{kalemdar1} ').bold = True
        p.add_run(f'kalendar kun; ')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph("d) shuningdek ushbu tashkilotda ishlagan har besh yili uchun qoʻshimcha davomiyligi ")
        p.add_run(f'{kalemdar2} ').bold = True
        p.add_run(f'kalendar kun (nomi) bo‘lgan haq to‘lanadigan yillik ta’til belgilanadi.')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY


        document.add_paragraph("14. Mehnat shartnomasining mehnat sharoitlari va unga haq to‘lash xususiyatlari, ijtimoiy himoya, imtiyozlar, kafolatlar va hokazolar bilan bog‘liq bo‘lgan boshqa shartlari.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("Tomonlarning manzillari va imzolari:").alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('')
        p.add_run("Xodim:                                                                             Ish beruvchi:")
        p = document.add_paragraph('')
        p.add_run(f"{xodim}                                              Daminov Mirzohid Islomovich")
        p = document.add_paragraph('')
        p.add_run("Manzil: __________________                                       Manzil:Buxoro shahar, Piridastgir koʻchasi 2-uy")
        p = document.add_paragraph('')
        p.add_run("__________________________                                     Tel: (65) 226-10-97")
        p = document.add_paragraph('')
        p.add_run("__________________________                                     _________________________")
        p = document.add_paragraph('')
        p.add_run("             (sana, imzo)                                                                                                                            (sana, muhr)").font.size = Pt(8)
        p = document.add_paragraph('')
        p.add_run("")
        document.add_paragraph("")
        p = document.add_paragraph("")
        p.add_run("     Mehnat shartnomasini bir nusxasini oldim").font.size = Pt(8)
        document.add_paragraph("               __________")


        document.add_heading()
        docx_file = BytesIO()
        document.save(docx_file)
        docx_file.seek(0)

        response = HttpResponse(docx_file.read())
        response['Content-Disposition'] = 'attachment; filename={}.docx'.format(soni.replace(" ", "_"))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response

    return render(request, 'index.html')

@login_required
def create_docx2(request):
    if request.method == 'POST':
        soni = request.POST.get('soni', '')
        sana = request.POST.get('sana', '')
        xodim = request.POST.get('xodim', '')
        korxona = request.POST.get('korxona', '')
        lavozim = request.POST.get('lavozim', '')
        asosiy = request.POST.get('asosiy', '')
        muayan = request.POST.get('muayan', '')
        boshlash = request.POST.get('boshlash', '')
        tomom = request.POST.get('tomom', '')
        sinov = request.POST.get('sinov', '')
        ish = request.POST.get('ish', '')
        kun = request.POST.get('kun', '')
        soat = request.POST.get('soat', '')
        razryad = request.POST.get('razryad', '')
        kalemdar1 = request.POST.get('kalemdar1', '')
        kalemdar2 = request.POST.get('kalemdar2', '')

        Profesir.objects.create(shartnoma_soni=soni, sana=sana, xodim=xodim,
                                korxona=korxona, lavozim=lavozim, asosiy_ish=asosiy, muddatli_mehnat=muayan,
                                ish_boshlash=boshlash, tomom_bolishi=tomom, sinov_muddati=sinov, ish_vaqti=ish,
                                kun=kun, soat=soat, razryad=razryad, kalemdar1=kalemdar1, kalemdar2=kalemdar2,
                                author=request.user)

        document = Document()

        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(0.5)

        style = document.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        document.add_paragraph(f'{soni}- SON MEHNAT ShARTNOMASI  ').alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph('')
        p = document.add_paragraph(f'               ')
        p.add_run(f'{sana}').bold = True
        p.add_run(f'                                                                                             ')
        p.add_run('Buxoro shahri').bold = True

        a = document.add_paragraph("1. Korxona (mulkchilikning barcha shakllaridagi tashkilot, muassasa, shu jumladan, ularning alohida tarkibiy bo‘linmalari ")
        a.add_run('Buxoro davlat pedagogika instituti rektori ').bold = True
        a.add_run(f'(direktori) ')
        a.add_run(f'Daminov Mirzohid Islomovich ').bold = True
        a.add_run('nomidan, keyingi o‘rinlarda «Ish beruvchi» deb ataladi va fuqaro ')
        a.add_run(f'{xodim} ').bold = True
        a.add_run('keyingi o‘rinlarda «Xodim» deb ataladi, mazkur shartnomani quyidagilar haqida tuzdik:')
        a.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('2. Xodim                                              ')
        p.add_run(f"{xodim}").bold = True
        p = document.add_paragraph('')
        p.add_run("(familiyasi, ismi va otasining ismi)").font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('')
        p.add_run(f'{korxona}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('')
        p.add_run('(korxona tarkibiy bo‘linmasi, sex, bo‘lim, uchastka, laboratoriya va shu kabilarning nomi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f"kasbi bo‘yicha ")
        p.add_run(f'{lavozim} ').bold = True
        p.add_run('lavozimiga ishga qabul qilinadi.')
        p = document.add_paragraph('')
        p.add_run('(ХALIKK bo‘yicha kasb, lavozimining to‘liq nomi, razryad, malaka toifasi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f'3. Shartnoma ')
        p.add_run(f'{asosiy} ').bold = True
        p.add_run('hisoblanadi. ')

        p = document.add_paragraph('')
        p.add_run('(asosiy ish, o‘rindoshlik, mavsumiy va boshqa)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f'4. Shartnoma muddati ')
        p.add_run(f'{muayan} ').bold = True

        p = document.add_paragraph('')
        p.add_run('(nomuayyan muddatga,uch yildan ko‘p bo‘lmagan muayyan muddatga (muddatli mehnat shartnomasi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph('5. Shartnoma bo‘yicha ishlash')
        p = document.add_paragraph(f'– boshlanishi ')
        p.add_run(f'{boshlash} ').bold = True

        p = document.add_paragraph(f'– тamom bo‘lishi ')
        p.add_run(f'{tomom}').bold = True

        p = document.add_paragraph(f'6. Sinov muddati ')
        p.add_run(f'{sinov}').bold = True


        p = document.add_paragraph('')
        p.add_run('(sinovsiz, sinov muddati)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('7. ')
        p.add_run('Xodimning majburiyatlari:').bold = True

        document.add_paragraph("a) O‘zbekiston Respublikasi Konstitutsiyasi, mehnat to‘g‘risidagi qonunchilik hamda boshqa qonunlar va qonunchilik hujjatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph('b) Lavozim yo‘riqnomasida belgilangan o‘z xizmat majburiyatlarini vijdonan bajarishi,').alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("d) Ish beruvchi, yuqori turuvchi organ va tashkilotlarning buyruqlari, topshiriqlari, qarorlari, ko‘rsatmalarini bajarish (bunday buyruq, topshiriq, qaror, ko‘rsatma yaqqol qonunchilik hujjatlariga zid bo‘lgan holatlar bundan mustasno);").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("e) mehnat va texnologiya intizomi (Ustav, Ichki mehnat tartibi qoidalari, Odob-axloq kodeksi va boshqa lokal hujjatlar)ga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("f) ish beruvchining qonuniy talablarini bajarish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("g) mehnatni muhofaza qilish, xavfsizlik texnikasi va ishlab chiqarish sanitariyasi talablariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("h) lavozim yo‘riqnomalariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("i) XALIKK, MM bilan nazarda tutilgan malaka majburiyatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("j) qonunchilik va boshqa normativ hujjatlarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("k) jamoa shartnomasi shartlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("l) institutning korrupsiyaga qarshi kurashish siyosati va sohaga oid boshqa normativ-huquqiy hujjatlarda qoʻyilgan maqsad va vazifalarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("m) oʻzi toʻgʻrisidagi shaxsiy ma’lumotlarning haqqoniyligini zimmasiga olish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("n) Ish beruvchining ruxsatisiz oshkor qilib bo‘lmaydigan (konfidensial) ma’lumotlarni o‘z ichiga oluvchi axborotlarni uchinchi shaxsga bermaslik, shuningdek, shaxsga oid ma’lumotlarni va xizmat doirasida foydalanilgan axborotlarni sir saqlanishiga rioya etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("o) institut professor-o‘qituvchilari va boshqa toifadagi xodimlar ish vaqti davomida o‘zaro munosabatlarda umume’tirof etilgan axloq qoidalariga rioya etish va boshqa xodim va talabalarga nisbatan xushmuomalada bo‘lish.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("p) xorijiy tillardan birida milliy yoki unga tenglashtirilgan mos darajadagi xalqaro sertifikatga ega bo‘lish hamda belgilangan muddatlarda yangilab borish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("q) xodim tomonidan qabul qilinadigan boshqa majburiyatlar, jumladan: ").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('I.O‘quv-uslubiy ishlar yoʻnalishida').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("-  yangi ma’ruza matnlari, mashq (masalalar) to‘plami, laboratoriya, ijodiy ishlari bo‘yicha uslubiy qo‘llanmalarni ishlab chiqish, yozish va nashr etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  tarqatma o‘quv materiallari, elektron o‘quv dasturlar va video mashg‘ulotlar (fan doirasida), shuningdek tegishli kompyuter dasturlaridan foydalangan holda taqdimot slaydlari, o‘qitish dasturlarini tayyorlash (yaratish);").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  darslik, o‘quv qo‘llanmalar yozish va nashr etishga tayyorlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  darslik, o‘quv qo‘llanmalarni to‘ldirilgan va o‘zgartirilgan holda qayta nashr etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  uslubiy qo‘llanma (ko‘rsatma, risola)lar tayyorlash va nashr etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  fanning o‘quv-uslubiy majmuasini ishlab chiqish va fanning elektron modul papkasini yaratish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  mavjud laboratoriya ishini yangilash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  yangi laboratoriya ishini tayyorlash va joriy etish, shuningdek virtual laboratoriya ishini tayyorlash va joriy etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  bakalavriat ta’lim yo‘nalishi (magistratura, doktorantura mutaxassisligi) bo‘yicha malaka talabi, namunaviy o‘quv reja, yangi fan uchun namunaviy o‘quv dasturlarini ishlab chiqish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  ishchi o‘quv reja, fan bo‘yicha ishchi o‘quv dasturini ishlab chiqish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  darslik, o‘quv qo‘llanmalar va boshqa materiallarni tarjima qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  fanlar bo‘yicha nazorat savollari (test, masalalar va boshqa), oraliq va yakuniy baholashlar uchun topshiriqlarni ishlab chiqish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('II. Ilmiy-tadqiqot ishlari yoʻnalishida:').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("-  ilmiy maqola, tezis tayyorlash va nashr etish, ixtiro (patent), dasturiy vositalar yaratish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  monografiya va risola nashr etish, axborot tahliliy materiallar tayyorlash, ijodiy ishlar yaratish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  shaxsiy (tematik) ijodiy asarlar katalogini nashr etish, yangi loyihalar yaratish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  darslik, monografiya, risola va o‘quv qo‘llanmalariga muharrirlik qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  oliy ta’lim muassasasi kengashi tomonidan tasdiqlangan mavzu bo‘yicha ilmiy-tadqiqot ishlarini bajarish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  darslik, o‘quv qo‘llanma, monografiya, risola, ilmiy maqola, dissertatsiya, avtoreferat va boshqa ishlarga taqriz yozish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  falsafa doktori (PhD), fan doktori (DsC) ilmiy darajasini olish uchun dissertatsiya tayyorlash va himoya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  ilmiy jurnallar tahririyatiga a’zo bo‘lish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  talabalarning ilmiy-tadqiqot va ilmiy-ijodiy ishlariga rahbarlik qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  talabalarning fan va ilmiy to‘garaklariga rahbarlik qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  ilmiy yoki uslubiy konferensiya va seminarlar tashkil etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  ilmiy yoki uslubiy kengashlarda ishtirok etish.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('III. «Ustoz-shogird» ishlari').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("-  talaba oliy ta’lim muassasasiga o‘qishga qabul qilinganda yangi sharoitga adaptatsiya jarayonini o‘tashga va oliy ta’lim muassasasida bo‘lgan chog‘ida vaqtini to‘g‘ri taqsimlashga yordam berish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  talabaning shaxsiy va akademik yutuqlarini rivojlantirish maqsadida uchrashuvlar o‘tkazish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  talabalar baholash tizimi qanday olib borilishini qay darajada anglaganini va qayerdan ta’lim va akademik sohada yordam olishi mumkin ekanligini nazorat qilish").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  o‘qish jarayonida shaxsiy va akademik muammolar yuzaga kelganda va ular ta’lim olish darajasiga ta’sir ko‘rsatsa, ularni hal qilishda ko‘maklashish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  akademik yoki kasbiy tanlov qilishda tavsiyalar berish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  ma’ruza materiallarni qaytarish va misollarni taqdim etgan holda ularni yoritish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  talabalarning mutaxassislik va tanlov fanlaridan olgan baholari va o‘zlashtirgan bilimlari ularning kvalifikatsion darajalariga ta’sirining muhimligi haqida tushunchani shakllantirish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  boshqa talaba va o‘qituvchilardan ushbu talaba haqida ma’lumot olishi;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("-  talabalarning ota-onalari bilan ishlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY



        p = document.add_paragraph('8. ')
        p.add_run('Ish beruvchining majburiyatlari:').bold = True

        document.add_paragraph("a) O‘zbekiston Respublikasi Konstitutsiyasi, mehnat to‘g‘risidagi qonunchilik hamda boshqa qonunlar va qonunchilik hujjatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("b) ishga qabul qilish chogʻida xodimni Ustav, Ichki mehnat tartib-qoidalari, Odob-axloq kodeksi, mehnatini tashkil etish, mehnatni muhofaza qilish va texnik xavfsizlikni ta’minlash qoidalari, Lavozim yo‘riqnomasi, jamoa shartnomasi va boshqa ichki lokal hujjatlar bilan tanishtirish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("d) mehnat va ishlab chiqarish intizomini ta’minlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("e) ish haqi va qonunchilik hujjatlarida nazarda tutilgan boshqa to‘lovlarni o‘z vaqtida hamda to‘liq hajmda to‘lash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("f) xizmat vazifalarini bajarishda zarur xavfsiz va samarali mehnat sharoitlarini yaratishni ta’minlash, uni kasbiy kompetensiyalarini oshirish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("g) ish joyini mehnatni muhofaza qilish va xavfsizlik texnikasi qoidalariga muvofiq jihozlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("h) qonunchilikka va boshqa normativ hujjatlarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("i) jamoa shartnomasi shartlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("j) ish beruvchi tomonidan qabul qilinadigan boshqa majburiyatlar:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("k)Ish beruvchi qonunchilik hujjatlarida belgilangan boshqa majburiyatlarni ham bajarishi mumkin.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("9.Mehnat shartnomasini tuzish va bekor qilish, shuningdek unga o‘zgartish va qo‘shimchalar kiritish qonunchilikda belgilangan tartibda «Yagona milliy mehnat tizimi» idoralararo dasturiy-apparat kompleksida ro‘yxatdan o‘tkazilishi shart.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph("10. Ish kuni rejimi ")
        p.add_run(f"{ish} ").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph('')
        p.add_run('(ish vaqtining normal davomiyligi; ish vaqtining qisqartirilgan davomiyligi; to‘liqsiz ish vaqti.').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f"11.Xodim uchun ")
        p.add_run(f'{kun} ').bold = True
        p.add_run('kunlik ish haftasi belgilanadi. Bunda ish vaqtining davomiyligi haftasiga ')
        p.add_run(f'{soat} ').bold = True
        p.add_run('soatgacha etib belgilanadi.')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("12. Mehnatga haq to‘lash.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("Xodimga quyidagicha haq to‘lash belgilanadi:")

        p = document.add_paragraph(f"a) ")
        p.add_run(f'{razryad} ').bold = True
        p.add_run(f"ming soʻm")
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('(to‘lov turi va uning aniq summadagi yoki YTS razryadi ko‘rsatilgan holdagi, yohud tushumdan olingan foizlardagi miqdori)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph("b) amaldagi qonunchilikka va normativ hujjatlarga muvofiq mehnat sharoitlari bilan bog‘liq bo‘lgan qo‘shimcha haq, ustama, kompensatsiyalar quyidagi miqdorlarda:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('        (qo‘shimcha haq, ustama, kompensatsiyalar nomi)                                            (ularning miqdori) ').font.size = Pt(10)

        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("v) jamoa shartnomasi shartlari bilan nazarda tutilgan, shuningdek, berilgan (mavjud) huquqlar va mablag‘lar doirasida rahbar tomonidan belgilanadigan qo‘shimcha haq, ustama, mukofot, taqdirlashlar va rag‘batlantiruvchi turdagi boshqa to‘lovlar;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('        (qo‘shimcha haq, ustama, kompensatsiyalar nomi)                                            (ularning miqdori) ').font.size = Pt(10)
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("13. Xodimga:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("a) asosiy ta’til (mehnat ta’tili) 21 kalendar kun;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph("b) jamoa shartnomasi boʻyicha qo‘shimcha ta’til ")
        p.add_run(f'{kalemdar1} ').bold = True
        p.add_run(f'kalendar kun; ')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph("d) shuningdek ushbu tashkilotda ishlagan har besh yili uchun qoʻshimcha davomiyligi ")
        p.add_run(f'{kalemdar2} ').bold = True
        p.add_run(f'kalendar kun (nomi) bo‘lgan haq to‘lanadigan yillik ta’til belgilanadi.')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY


        document.add_paragraph("14. Mehnat shartnomasining mehnat sharoitlari va unga haq to‘lash xususiyatlari, ijtimoiy himoya, imtiyozlar, kafolatlar va hokazolar bilan bog‘liq bo‘lgan boshqa shartlari.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("Tomonlarning manzillari va imzolari:").alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('')
        p.add_run("Xodim:                                                                             Ish beruvchi:")
        p = document.add_paragraph('')
        p.add_run(f"{xodim}                                              Daminov Mirzohid Islomovich ")
        p = document.add_paragraph('')
        p.add_run(
            "Manzil: __________________                                       Manzil:Buxoro shahar, Piridastgir koʻchasi 2-uy")
        p = document.add_paragraph('')
        p.add_run("__________________________                                     Tel: (65) 226-10-97")
        p = document.add_paragraph('')
        p.add_run("__________________________                                     _________________________")
        p = document.add_paragraph('')
        p.add_run("             (sana, imzo)                                                                                                                            (sana, muhr)").font.size = Pt(8)
        p = document.add_paragraph('')
        p.add_run("")
        document.add_paragraph("")
        p = document.add_paragraph("")
        p.add_run("     Mehnat shartnomasini bir nusxasini oldim").font.size = Pt(8)
        document.add_paragraph("               __________")


        document.add_heading()
        docx_file = BytesIO()
        document.save(docx_file)
        docx_file.seek(0)

        response = HttpResponse(docx_file.read())
        response['Content-Disposition'] = 'attachment; filename={}.docx'.format(soni.replace(" ", "_"))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response

    return render(request, 'index2.html')

@login_required
def create_docx3(request):
    if request.method == 'POST':
        soni = request.POST.get('soni', '')
        sana = request.POST.get('sana', '')
        xodim = request.POST.get('xodim', '')
        korxona = request.POST.get('korxona', '')
        lavozim = request.POST.get('lavozim', '')
        hissa = request.POST.get('hissa', '')
        asosiy = request.POST.get('asosiy', '')
        muayan = request.POST.get('muayan', '')
        boshlash = request.POST.get('boshlash', '')
        tomom = request.POST.get('tomom', '')
        sinov = request.POST.get('sinov', '')
        ish = request.POST.get('ish', '')
        kun = request.POST.get('kun', '')
        soat = request.POST.get('soat', '')
        razryad = request.POST.get('razryad', '')
        kalemdar = request.POST.get('kalemdar', '')
        kalemdar1 = request.POST.get('kalemdar1', '')
        kalemdar2 = request.POST.get('kalemdar2', '')

        Orindosh.objects.create(shartnoma_soni=soni, sana=sana, xodim=xodim,
                                korxona=korxona, lavozim=lavozim, hissa=hissa, orindoshlik=asosiy, muddatli_mehnat=muayan,
                                ish_boshlash=boshlash, tomom_bolishi=tomom, sinov_muddati=sinov, ish_vaqti=ish,
                                kun=kun, soat=soat, razryad=razryad,kalemdar=kalemdar, kalemdar1=kalemdar1, kalemdar2=kalemdar2,
                                author=request.user)

        document = Document()

        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(0.5)

        style = document.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

        document.add_paragraph(f'{soni}- SON O‘RINDOSHLIK ASOSIDA ISHLASH TO‘G‘RISIDA').alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('MEHNAT SHARTNOMASI').alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph('')
        p = document.add_paragraph(f'               ')
        p.add_run(f'{sana}').bold = True
        p.add_run(f'                                                                                             ')
        p.add_run('Buxoro shahri').bold = True

        a = document.add_paragraph("1. Korxona (mulkchilikning barcha shakllaridagi tashkilot, muassasa, shu jumladan, ularning alohida tarkibiy bo‘linmalari ")
        a.add_run('Buxoro davlat pedagogika instituti rektori ').bold = True
        a.add_run(f'(direktori) ')
        a.add_run(f'Daminov Mirzohid Islomovich ').bold = True
        a.add_run('nomidan, keyingi o‘rinlarda «Ish beruvchi» deb ataladi va fuqaro ')
        a.add_run(f'{xodim} ').bold = True
        a.add_run('keyingi o‘rinlarda «Xodim» deb ataladi, mazkur shartnomani quyidagilar haqida tuzdik:')
        a.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('2. Xodim                                              ')
        p.add_run(f"{xodim}").bold = True
        p = document.add_paragraph('')
        p.add_run("(familiyasi, ismi va otasining ismi)").font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('')
        p.add_run(f'{korxona}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('')
        p.add_run('(korxona tarkibiy bo‘linmasi, sex, bo‘lim, uchastka, laboratoriya va shu kabilarning nomi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f"3. kasbi bo‘yicha ")
        p.add_run(f'{lavozim} ').bold = True
        p.add_run('lavozimiga ')
        p.add_run(f'{hissa} ').bold = True
        p.add_run('hissa ')
        p.add_run(f'{asosiy} ').bold = True
        p.add_run('o‘rindoshlik asosida ishga qabul qilinadi.')

        p = document.add_paragraph('')
        p.add_run('(ХALIKK bo‘yicha kasb, lavozimining to‘liq nomi, razryad, malaka toifasi)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f'4. Shartnoma muddati ')
        p.add_run(f'{muayan} ').bold = True

        p = document.add_paragraph('')
        p.add_run('( asosiy xodim ishga qabul qilingunga qadar, o’quv yuklamalarini bajargunga qadar, asosiy xodim ishga chiqqunga qadar)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph('5. Shartnoma bo‘yicha ishlash')
        p = document.add_paragraph(f'– boshlanishi ')
        p.add_run(f'{boshlash} ').bold = True

        p = document.add_paragraph(f'– тamom bo‘lishi ')
        p.add_run(f'{tomom}').bold = True

        p = document.add_paragraph(f'6. Sinov muddati ')
        p.add_run(f'{sinov}').bold = True


        p = document.add_paragraph('')
        p.add_run('(sinovsiz, sinov muddati)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph('7. ')
        p.add_run('Xodimning majburiyatlari:').bold = True

        document.add_paragraph("a) O‘zbekiston Respublikasi Konstitutsiyasi, mehnat to‘g‘risidagi qonunchilik hamda boshqa qonunlar va qonunchilik hujjatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph('b) Lavozim yo‘riqnomasida belgilangan o‘z xizmat majburiyatlarini vijdonan bajarishi,').alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("d) Ish beruvchi, yuqori turuvchi organ va tashkilotlarning buyruqlari, topshiriqlari, qarorlari, ko‘rsatmalarini bajarish (bunday buyruq, topshiriq, qaror, ko‘rsatma yaqqol qonunchilik hujjatlariga zid bo‘lgan holatlar bundan mustasno);").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("e) mehnat va texnologiya intizomi (Ustav, Ichki mehnat tartibi qoidalari, Odob-axloq kodeksi va boshqa lokal hujjatlar)ga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("f) ish beruvchining qonuniy talablarini bajarish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("g) mehnatni muhofaza qilish, xavfsizlik texnikasi va ishlab chiqarish sanitariyasi talablariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("h) lavozim yo‘riqnomalariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("i) XALIKK, MM bilan nazarda tutilgan malaka majburiyatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("j) qonunchilik va boshqa normativ hujjatlarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("k) jamoa shartnomasi shartlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("l) institutning korrupsiyaga qarshi kurashish siyosati va sohaga oid boshqa normativ-huquqiy hujjatlarda qoʻyilgan maqsad va vazifalarga rioya qilish; Shuningdek:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - xodim atrofdagi shaxslar tomonidan korrupsiyaga oid huquqbuzarlikni sodir etish yoki unda qatnashishga tayyorgarlik ko‘rish sifatida baholanishi mumkin bo‘lgan xulq-atvordan o‘zini tiyishi shart;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - Institutning boshqa xodimlari tomonidan sodir etilgan yoki sodir etilishi rejalashtirilayotgan korrupsiyaga oid huquqbuzarliklar haqidagi dalillarni aniqlasa, ushbu ma’lumotlarni tegishli vakolatli organlarga xabar berishi lozim;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - xodim O‘zbekiston Respublikasi qonun hujjatlari va institutning ichki hujjatlariga muvofiq manfaatlar to‘qnashuvi kelib chiqishini oldini olish choralarini ko‘rish, manfaatlar to‘qnashuvi yoki uning kelib chiqish ehtimoli haqida o‘ziga ma’lum bo‘lishi bilan darhol ish beruvchini yoki uning o‘rnini bosadigan shaxslarni xabardor qilish shart;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("m) oʻzi toʻgʻrisidagi shaxsiy ma’lumotlarning haqqoniyligini zimmasiga olish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("n) Ish beruvchining ruxsatisiz oshkor qilib bo‘lmaydigan (konfidensial) ma’lumotlarni o‘z ichiga oluvchi axborotlarni uchinchi shaxsga bermaslik, shuningdek, shaxsga oid ma’lumotlarni va xizmat doirasida foydalanilgan axborotlarni sir saqlanishiga rioya etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("o) institut professor-o‘qituvchilari va boshqa toifadagi xodimlar ish vaqti davomida o‘zaro munosabatlarda umume’tirof etilgan axloq qoidalariga rioya etish va boshqa xodim va talabalarga nisbatan xushmuomalada bo‘lish.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("p) axborot xavfsizligini ta’minlash, jumladan, internet tarmoqlaridan faqat soha bo‘yicha foydalanish, ijtimoiy tarmoqlardan ish faoliyatidan chetga chiqqan holda foydalanmaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("q) xodim o‘z xizmat vazifalarini bajarayotganda korrupsiyaviy hatti-harakatlarda qatnashmaslik, shu jumladan g‘ayriqonuniy imtiyozlarni olish yoki boshqacha tarzda g‘ayriqonuniy maqsadlarda o‘zining yoki boshqa shaxslarning harakatlari yoki qarorlariga  ta'sir qilish uchun pul mablag‘lari yoki boshqa qimmatliklar ko‘rinishida: ").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - pora taklif qilmaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - va’da bermaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - pul to‘lamaslik, tovlamachilik qilmaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - bevosita yoki bilvosita pora olishga rozilik bermaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - pora olmaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - biron-bir shaxs foydasiga yoki biron-bir shaxs xizmatlaridan foydalanmaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - mol-mulk yoki mulkiy huquqlarni olmaslik;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - o‘z vakolatlarini suiiste’mol qilmaslik majburiyatini oladi;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - lavozim vazifalarini halol bajarish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - ommaviy jamoat ishlarida bevosita ishtirok etish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("  - o‘zining xizmat vazifalarini o‘z vaqtida, to‘liq, halol va vijdonan bajarish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("r) Xodim qabul qilgan boshqa majbuyiyat _____________________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY


        p = document.add_paragraph('8. ')
        p.add_run('Ish beruvchining majburiyatlari:').bold = True

        document.add_paragraph("a) O‘zbekiston Respublikasi Konstitutsiyasi, mehnat to‘g‘risidagi qonunchilik hamda boshqa qonunlar va qonunchilik hujjatlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("b) ishga qabul qilish chogʻida xodimni Ustav, Ichki mehnat tartib-qoidalari, Odob-axloq kodeksi, mehnatini tashkil etish, mehnatni muhofaza qilish va texnik xavfsizlikni ta’minlash qoidalari, Lavozim yo‘riqnomasi, jamoa shartnomasi va boshqa ichki lokal hujjatlar bilan tanishtirish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("d) mehnat va ishlab chiqarish intizomini ta’minlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("e) ish haqi va qonunchilik hujjatlarida nazarda tutilgan boshqa to‘lovlarni o‘z vaqtida hamda to‘liq hajmda to‘lash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("f) xizmat vazifalarini bajarishda zarur xavfsiz va samarali mehnat sharoitlarini yaratishni ta’minlash, uni kasbiy kompetensiyalarini oshirish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("g) ish joyini mehnatni muhofaza qilish va xavfsizlik texnikasi qoidalariga muvofiq jihozlash;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("h) qonunchilikka va boshqa normativ hujjatlarga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("i) jamoa shartnomasi shartlariga rioya qilish;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("j) ish beruvchi tomonidan qabul qilinadigan boshqa majburiyatlar:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("k)Ish beruvchi qonunchilik hujjatlarida belgilangan boshqa majburiyatlarni ham bajarishi mumkin.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("9.Mehnat shartnomasini tuzish va bekor qilish, shuningdek unga o‘zgartish va qo‘shimchalar kiritish qonunchilikda belgilangan tartibda «Yagona milliy mehnat tizimi» idoralararo dasturiy-apparat kompleksida ro‘yxatdan o‘tkazilishi shart.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph("10. Ish kuni rejimi ")
        p.add_run(f"{ish} ").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph('')
        p.add_run('(ish vaqtining normal davomiyligi; ish vaqtining qisqartirilgan davomiyligi; to‘liqsiz ish vaqti.').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = document.add_paragraph(f"11.Xodim uchun ")
        p.add_run(f'{kun} ').bold = True
        p.add_run('kunlik ish haftasi belgilanadi. Bunda ish vaqtining davomiyligi haftasiga ')
        p.add_run(f'{soat} ').bold = True
        p.add_run('soatgacha etib belgilanadi.')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("12. Mehnatga haq to‘lash.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("Xodimga quyidagicha haq to‘lash belgilanadi:")

        p = document.add_paragraph(f"a) ")
        p.add_run(f'Daminov Mirzohid Islomovich ').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('(to‘lov turi va uning aniq summadagi yoki YTS razryadi ko‘rsatilgan holdagi, yohud tushumdan olingan foizlardagi miqdori)').font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph("b) amaldagi qonunchilikka va normativ hujjatlarga muvofiq mehnat sharoitlari bilan bog‘liq bo‘lgan qo‘shimcha haq, ustama, kompensatsiyalar quyidagi miqdorlarda:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('        (qo‘shimcha haq, ustama, kompensatsiyalar nomi)                                            (ularning miqdori) ').font.size = Pt(10)

        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("d) jamoa shartnomasi shartlari bilan nazarda tutilgan, shuningdek, berilgan (mavjud) huquqlar va mablag‘lar doirasida rahbar tomonidan belgilanadigan qo‘shimcha haq, ustama, mukofot, taqdirlashlar va rag‘batlantiruvchi turdagi boshqa to‘lovlar;").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph('')
        p.add_run('        (qo‘shimcha haq, ustama, kompensatsiyalar nomi)                                            (ularning miqdori) ').font.size = Pt(10)
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("____________________________________________ _______________________________________").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        document.add_paragraph("13. Xodimga:").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        p = document.add_paragraph(f"a) asosiy ta’til (mehnat ta’tili) ")
        p.add_run(f"{kalemdar} ").bold = True
        p.add_run('kalendar kun;')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph("b) jamoa shartnomasi boʻyicha qo‘shimcha ta’til ")
        p.add_run(f'{kalemdar1} ').bold = True
        p.add_run(f'kalendar kun; ')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY

        p = document.add_paragraph("d) shuningdek ushbu tashkilotda ishlagan har besh yili uchun qoʻshimcha davomiyligi ")
        p.add_run(f'{kalemdar2} ').bold = True
        p.add_run(f'kalendar kun (nomi) bo‘lgan haq to‘lanadigan yillik ta’til belgilanadi.')
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY


        document.add_paragraph("14. Mehnat shartnomasining mehnat sharoitlari va unga haq to‘lash xususiyatlari, ijtimoiy himoya, imtiyozlar, kafolatlar va hokazolar bilan bog‘liq bo‘lgan boshqa shartlari.").alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
        document.add_paragraph("Tomonlarning manzillari va imzolari:").alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = document.add_paragraph('')
        p.add_run("Xodim:                                                                             Ish beruvchi:")
        p = document.add_paragraph('')
        p.add_run(f"{xodim}                                              {rektor}")
        p = document.add_paragraph('')
        p.add_run(
            "Manzil: __________________                                       Manzil:Buxoro shahar, Piridastgir koʻchasi 2-uy")
        p = document.add_paragraph('')
        p.add_run("__________________________                                     Tel: (65) 226-10-97")
        p = document.add_paragraph('')
        p.add_run("__________________________                                     _________________________")
        p = document.add_paragraph('')
        p.add_run("             (sana, imzo)                                                                                                                            (sana, muhr)").font.size = Pt(8)
        p = document.add_paragraph('')
        p.add_run("")
        document.add_paragraph("")
        p = document.add_paragraph("")
        p.add_run("     Mehnat shartnomasini bir nusxasini oldim").font.size = Pt(8)
        document.add_paragraph("               __________")


        document.add_heading()
        docx_file = BytesIO()
        document.save(docx_file)
        docx_file.seek(0)

        response = HttpResponse(docx_file.read())
        response['Content-Disposition'] = 'attachment; filename={}.docx'.format(soni.replace(" ", "_"))
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response

    return render(request, 'index3.html')